<USER_REQUEST>
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

qs=[]

def add(q, opts, ans, exp, topic, level="NEET", typ="Single Correct MCQ"):
    qs.append(dict(q=q,opts=opts,ans=ans,exp=exp,topic=topic,level=level,type=typ))

# 1-30 Mendelian inheritance
data = [
("Mendel selected garden pea for his experiments mainly because it:",["Has only asexual reproduction","Shows many clear contrasting traits and permits controlled pollination","Has an exceptionally long generation time","Produces only one seed per generation"],"B","Pea plants have distinct contrasting characters, a short generation time and can be selfed or cross-pollinated.","Mendel's experimental material"),
("The number of pairs of contrasting characters studied by Mendel was:",["5","6","7","8"],"C","Mendel studied seven pairs of contrasting traits in pea.","Mendel's experiments"),
("A true-breeding organism is one that:",["Always produces hybrids","Produces offspring resembling the parent for a trait after repeated selfing","Is necessarily homozygous dominant","Cannot reproduce sexually"],"B","True-breeding lines show stable trait inheritance after continuous self-pollination.","True breeding"),
("In a monohybrid cross TT × tt, all F1 offspring are:",["TT","Tt","tt","1 TT:1 tt"],"B","Each offspring receives T from one parent and t from the other.","Monohybrid cross"),
("Selfing the F1 of TT × tt gives an F2 phenotypic ratio of:",["1:1","3:1","1:2:1","9:3:3:1"],"B","Complete dominance produces three dominant phenotypes for every one recessive phenotype.","Monohybrid cross"),
("The F2 genotypic ratio in a monohybrid cross is:",["3:1","1:1","1:2:1","9:3:3:1"],"C","Segregation of T and t from Tt parents gives TT:Tt:tt = 1:2:1.","Law of segregation"),
("A test cross is performed between an individual of unknown dominant genotype and:",["A homozygous dominant individual","A heterozygous individual","A homozygous recessive individual","Any F1 hybrid"],"C","The recessive tester reveals the gametes produced by the unknown genotype.","Test cross"),
("A tall pea plant gives tall and dwarf progeny in a test cross. Its genotype is:",["TT","Tt","tt","Cannot be determined"],"B","Production of recessive offspring proves that the tall parent carries t.","Test cross"),
("A tall plant produces only tall offspring when crossed with a dwarf plant. The most likely genotype of the tall plant is:",["TT","Tt","tt","Tt or tt"],"A","TT × tt produces only Tt tall offspring.","Test cross"),
("The law of segregation states that:",["Alleles of different genes always assort together","Two alleles of a gene separate during gamete formation","Dominant alleles destroy recessive alleles","Genes blend in the offspring"],"B","The two alleles remain discrete and segregate into different gametes.","Law of segregation"),
("Mendel's law of dominance explains why:",["Both alleles enter the same gamete","Only one of two contrasting traits appears in F1","Recessive alleles disappear permanently","Genes on the same chromosome assort independently"],"B","In a heterozygote, the dominant allele expresses while the recessive allele is masked.","Law of dominance"),
("Which observation most directly supports the particulate nature of inheritance?",["F1 plants are all tall","The recessive trait reappears unchanged in F2","Plants produce many seeds","Pea flowers are bisexual"],"B","Reappearance of the recessive trait shows that hereditary factors do not blend.","Particulate inheritance"),
("A heterozygous tall pea produces gametes in the ratio:",["Only T","Only t","1 T:1 t","3 T:1 t"],"C","Alleles segregate equally during gamete formation.","Gamete formation"),
("The probability of obtaining a dwarf offspring from Tt × Tt is:",["1/2","1/4","3/4","1"],"B","Only tt is dwarf, with probability 1/4.","Probability in monohybrid cross"),
("The probability of a homozygous tall offspring from Tt × Tt is:",["1/4","1/2","3/4","1"],"A","TT occurs in one of four zygotic combinations.","Probability in monohybrid cross"),
("Among tall F2 plants from Tt × Tt, the fraction that is heterozygous is:",["1/3","1/2","2/3","3/4"],"C","Tall genotypes are TT and Tt in a 1:2 ratio; 2 of 3 tall plants are Tt.","Conditional probability"),
("A back cross is a cross of an F1 hybrid with:",["An unrelated individual","Either of its parents","Only the recessive parent","Another F1 only"],"B","Crossing F1 with either parental genotype is a back cross.","Back cross"),
("Every test cross is a back cross, but every back cross is not a test cross because:",["A test cross uses only the dominant parent","A test cross specifically uses the homozygous recessive parent","A back cross never uses a parent","Test crosses involve two traits"],"B","A back cross with the dominant parent is not a test cross.","Back cross and test cross"),
("In a dihybrid cross RrYy × RrYy, the F2 phenotypic ratio under independent assortment is:",["3:1","1:2:1","9:3:3:1","1:1:1:1"],"C","Independent assortment of two completely dominant gene pairs produces 9:3:3:1.","Dihybrid cross"),
("The number of different gamete types produced by RrYy is:",["2","4","8","16"],"B","The gametes are RY, Ry, rY and ry.","Dihybrid gametes"),
("The number of gamete types produced by AaBbCc is:",["3","6","8","16"],"C","For n heterozygous independently assorting loci, gamete types = 2^n = 8.","Gamete types"),
("The probability of aabb in AaBb × AaBb is:",["1/4","1/8","1/16","3/16"],"C","P(aa)=1/4 and P(bb)=1/4; multiply to obtain 1/16.","Dihybrid probability"),
("The probability of A_B_ offspring in AaBb × AaBb is:",["3/16","6/16","9/16","12/16"],"C","Dominant phenotype probability at each locus is 3/4; product is 9/16.","Dihybrid probability"),
("In a dihybrid test cross AaBb × aabb, independent assortment gives:",["3:1","1:2:1","9:3:3:1","1:1:1:1"],"D","The heterozygote forms four gamete types equally.","Dihybrid test cross"),
("The law of independent assortment is best demonstrated by:",["A monohybrid F2 ratio","A dihybrid F2 ratio","A test cross involving one gene","Incomplete dominance"],"B","The 9:3:3:1 ratio reflects independent assortment of two gene pairs.","Independent assortment"),
("In AaBbCc × aabbcc, the probability of aabbcc is:",["1/4","1/8","1/16","1/64"],"B","The triple heterozygote produces abc gametes with probability 1/8.","Multihybrid test cross"),
("How many distinct genotypes are possible in the offspring of AaBb × AaBb?",["4","8","9","16"],"C","Each locus has three genotypes, giving 3 × 3 = 9.","Dihybrid genotypes"),
("How many phenotypes occur in AaBb × AaBb with complete dominance and independent assortment?",["2","3","4","9"],"C","The four phenotypic classes are A_B_, A_bb, aaB_ and aabb.","Dihybrid phenotypes"),
("The probability of an offspring heterozygous at both loci from AaBb × AaBb is:",["1/4","1/8","1/16","1/2"],"A","P(Aa)=1/2 and P(Bb)=1/2, so P(AaBb)=1/4.","Dihybrid probability"),
("If genes A and B assort independently, the probability of gamete Ab from AaBb is:",["1/2","1/4","1/8","1/16"],"B","Each allele choice has probability 1/2; their product is 1/4.","Independent assortment")
]
for x in data: add(*x)

# 31-55 deviations, allelism, polygenic
data2=[
("In incomplete dominance, the F1 phenotype is:",["Identical to the dominant parent","Intermediate between the two parents","Identical to the recessive parent","Absent"],"B","Neither allele is completely dominant, so the heterozygote is intermediate.","Incomplete dominance"),
("In Snapdragon, red-flowered RR crossed with white rr produces F1 flowers that are:",["Red","White","Pink","Red and white spotted"],"C","Rr has an intermediate pink phenotype.","Incomplete dominance"),
("The F2 phenotypic ratio in incomplete dominance is:",["3:1","1:2:1","9:3:3:1","1:1"],"B","Each genotype has a distinct phenotype, so phenotype and genotype ratios coincide.","Incomplete dominance"),
("Incomplete dominance does not violate the law of segregation because:",["Alleles blend permanently","The two alleles still separate during gamete formation","Only one allele enters all gametes","Heterozygotes produce no gametes"],"B","Expression is intermediate, but alleles remain discrete and segregate.","Incomplete dominance"),
("Codominance differs from incomplete dominance because in codominance:",["The heterozygote is intermediate","Both alleles express fully in the heterozygote","The recessive allele disappears","Only one allele is inherited"],"B","Both allelic products are simultaneously expressed.","Codominance"),
("The human AB blood group is an example of:",["Incomplete dominance","Codominance","Polygenic inheritance","Pleiotropy"],"B","IA and IB are both expressed in IAIB individuals.","ABO blood group"),
("The ABO blood group system is controlled by:",["Two alleles","Three alleles in the population","Four alleles in each person","One allele with environmental effects"],"B","IA, IB and i constitute multiple alleles, though each individual carries only two.","Multiple alleles"),
("The genotype of a person with blood group O is:",["IAIA","IBIB","IAIB","ii"],"D","The O phenotype occurs only in homozygous ii individuals.","ABO blood group"),
("A person with blood group AB can produce gametes carrying:",["Only IA","Only IB","IA or IB","IA, IB or i"],"C","An IAIB individual segregates IA and IB.","ABO blood group"),
("Parents with blood groups A and B have a child with blood group O. The parental genotypes must be:",["IAIA and IBIB","IAi and IBi","IAIB and ii","IAIA and IBi"],"B","An O child receives i from each parent.","ABO inheritance"),
("A cross IAi × IBi can produce children with:",["Only A and B groups","Only AB and O groups","A, B, AB and O groups","Only O group"],"C","The four genotypes IAIB, IAi, IBi and ii are equally possible.","ABO inheritance"),
("An AB parent and an O parent can have offspring with blood groups:",["A or B only","AB or O only","A, B, AB or O","O only"],"A","IAIB × ii gives IAi or IBi.","ABO inheritance"),
("Multiple alleles means that:",["An individual possesses more than two alleles of a gene","More than two alternative alleles exist in a population","A gene controls several traits","Many genes control one trait"],"B","A diploid individual still carries only two alleles at a locus.","Multiple alleles"),
("Pleiotropy refers to:",["One trait controlled by many genes","One gene influencing more than one phenotypic trait","Two genes occupying the same locus","Expression of both alleles"],"B","A pleiotropic gene has multiple phenotypic effects.","Pleiotropy"),
("Phenylketonuria illustrates pleiotropy because a single defective gene can affect:",["Only skin colour","Several traits through altered phenylalanine metabolism","Only chromosome number","Only blood group"],"B","The metabolic defect has multiple downstream phenotypic effects.","Pleiotropy"),
("Polygenic inheritance involves:",["One gene with many alleles","Many genes contributing additively to one trait","One gene controlling many traits","Only sex-linked genes"],"B","Multiple loci contribute to continuous phenotypic variation.","Polygenic inheritance"),
("Human skin colour is described in NCERT as an example of:",["Codominance","Polygenic inheritance","Multiple allelism","Complete dominance"],"B","Several gene pairs contribute additively to pigmentation.","Polygenic inheritance"),
("If three independently assorting gene pairs contribute additively to skin colour, the number of phenotypic classes expected is:",["3","6","7","8"],"C","For n additive gene pairs, phenotypic classes = 2n+1 = 7.","Polygenic inheritance"),
("With three additive gene pairs, the probability of the darkest phenotype in F2 is:",["1/8","1/16","1/32","1/64"],"D","The darkest phenotype requires all six contributing alleles; probability is (1/4)^3=1/64.","Polygenic inheritance"),
("Continuous variation in a polygenic trait arises mainly because:",["Alleles blend","Several genes have cumulative effects, often modified by environment","Chromosomes fail to segregate","Only one genotype occurs"],"B","Numerous additive genotypes create a graded distribution.","Polygenic inheritance"),
("Which pair is correctly matched?",["ABO blood group - incomplete dominance","Snapdragon flower colour - codominance","Human skin colour - polygenic inheritance","Phenylketonuria - multiple allelism"],"C","Human skin colour is a standard polygenic example.","Inheritance patterns"),
("In incomplete dominance, a test cross of the F1 with the recessive parent gives:",["All intermediate","1 intermediate:1 recessive","3 intermediate:1 recessive","1 dominant:1 recessive"],"B","Rr × rr gives Rr and rr equally.","Incomplete dominance"),
("A pink Snapdragon crossed with another pink Snapdragon produces red, pink and white in the ratio:",["1:1:1","1:2:1","3:1","9:3:3:1"],"B","Rr × Rr gives RR:Rr:rr = 1:2:1.","Incomplete dominance"),
("Which statement about dominance is correct?",["Dominance means the dominant allele is more frequent","Dominance describes the phenotype of a heterozygote","Dominant alleles are always beneficial","Recessive alleles are never expressed"],"B","Dominance is a relationship between alleles in heterozygous condition.","Dominance concept"),
("A trait controlled by one gene but strongly influenced by environment is not necessarily:",["Heritable","Phenotypically variable","Polygenic","Expressed"],"C","Environmental influence alone does not establish polygenic control.","Gene-environment interaction")
]
for x in data2: add(*x)

# 56-80 chromosome theory/linkage
data3=[
("The chromosomal theory of inheritance was proposed by:",["Mendel and Darwin","Sutton and Boveri","Watson and Crick","Morgan and Sturtevant"],"B","Sutton and Boveri related chromosome behaviour to Mendelian factors.","Chromosomal theory"),
("The parallel between genes and chromosomes includes that both:",["Occur only in gametes","Exist in pairs in diploid cells and segregate during gamete formation","Are made only of protein","Blend during fertilisation"],"B","Homologous chromosomes and alleles show paired occurrence and segregation.","Chromosomal theory"),
("Thomas Hunt Morgan used which organism extensively for genetic studies?",["Pisum sativum","Drosophila melanogaster","Neurospora crassa","Escherichia coli"],"B","Fruit flies have short generations and easily observed variations.","Morgan's experiments"),
("A major advantage of Drosophila for genetic experiments is:",["Only one chromosome","Long generation time","Many offspring and easily distinguishable sexes","Exclusive asexual reproduction"],"C","Drosophila breeds quickly and produces many progeny with visible traits.","Drosophila"),
("Genes located on the same chromosome tend to be inherited together due to:",["Dominance","Linkage","Mutation","Polyploidy"],"B","Physical association on a chromosome causes linkage.","Linkage"),
("Recombination between linked genes results mainly from:",["Independent assortment of homologues only","Crossing over between homologous chromosomes","DNA replication","Nondisjunction"],"B","Crossing over exchanges segments and creates recombinant allele combinations.","Recombination"),
("The frequency of recombination between two genes is generally proportional to:",["Their dominance","The distance between them on a chromosome","Their allele frequency","Chromosome number"],"B","Genes farther apart have a greater chance of a crossover between them.","Gene mapping"),
("One map unit is equivalent to:",["1% mutation frequency","1% recombination frequency","10% recombination frequency","One nucleotide"],"B","A centimorgan corresponds to one percent recombinant progeny.","Gene mapping"),
("If two genes show 12% recombination, their map distance is approximately:",["6 cM","12 cM","24 cM","88 cM"],"B","For short intervals, recombination percentage approximates map distance.","Gene mapping"),
("Complete linkage produces:",["Only recombinant combinations","Only parental combinations","A 9:3:3:1 ratio","Equal parental and recombinant classes"],"B","Without crossing over, linked alleles remain in parental combinations.","Complete linkage"),
("In a test cross AB/ab × ab/ab with complete linkage, offspring are expected as:",["1 AB/ab :1 ab/ab","1 Ab/ab :1 aB/ab","1:1:1:1","3:1"],"A","The heterozygote produces only AB and ab gametes.","Complete linkage"),
("In the same test cross, if genes assort independently, offspring ratio is:",["1:1","3:1","1:1:1:1","9:3:3:1"],"C","AB/ab produces four gamete classes equally when unlinked.","Independent assortment"),
("When linked dominant alleles are on one chromosome and recessive alleles on the homologous chromosome, the arrangement is:",["Repulsion","Coupling or cis","Translocation","Inversion"],"B","AB/ab is the coupling or cis arrangement.","Linkage phase"),
("The arrangement Ab/aB is called:",["Coupling","Cis","Repulsion or trans","Complete dominance"],"C","Each homolog carries one dominant and one recessive allele.","Linkage phase"),
("Morgan observed that the proportion of parental combinations was greater than recombinants because:",["Alleles blended","Genes were linked on the same chromosome","Mutation eliminated recombinants","Gametes were diploid"],"B","Linkage preserves parental allele combinations.","Morgan's linkage studies"),
("Which statement about strongly linked genes is correct?",["They always show 50% recombination","They are usually close together and show low recombination","They occupy identical loci","They are necessarily alleles"],"B","Close genes are less likely to be separated by crossing over.","Linkage"),
("A recombination frequency of 50% usually indicates that genes are:",["Very tightly linked","Effectively unlinked or very far apart","Allelic","Completely dominant"],"B","Independent assortment yields a maximum observable recombination of 50%.","Recombination"),
("The first genetic maps were prepared using recombination frequencies by:",["Sutton","Sturtevant","Mendel","de Vries"],"B","Sturtevant used recombination data to map Drosophila genes.","Gene mapping"),
("Crossing over occurs between:",["Sister chromatids only","Non-sister chromatids of homologous chromosomes","Non-homologous chromosomes only","Centromeres of homologues"],"B","Exchange between non-sister chromatids creates recombinants.","Crossing over"),
("The stage of meiosis most directly associated with crossing over is:",["Leptotene","Pachytene","Anaphase II","Telophase I"],"B","Crossing over occurs during pachytene of prophase I.","Crossing over"),
("Chiasmata become clearly visible during:",["Zygotene","Pachytene","Diplotene","Metaphase II"],"C","Chiasmata are visible when homologues begin separating in diplotene.","Meiosis and linkage"),
("If parental offspring are 420 and 430 while recombinants are 70 and 80, recombination frequency is:",["7.5%","15%","30%","85%"],"B","Recombinants=150 of 1000 total, so frequency=15%.","Recombination calculation","JEE Main"),
("Two linked genes produce 18% recombinants in a test cross. The total parental percentage is:",["18%","36%","64%","82%"],"D","Parental classes together constitute 100-18=82%.","Recombination calculation"),
("In a test cross with 20% recombination, each recombinant class is expected to constitute:",["5%","10%","20%","40%"],"B","The two reciprocal recombinant classes share 20% equally.","Recombination calculation"),
("Linkage and recombination are best regarded as:",["Unrelated processes","Opposing consequences of gene location and crossing over","Identical terms","Forms of dominance"],"B","Linkage preserves combinations, whereas crossing over generates new combinations.","Linkage and recombination")
]
for x in data3: add(*x)

# 81-105 sex determination, pedigree
data4=[
("In humans, females are homogametic because they produce:",["X and Y eggs","Only X-bearing eggs","Only Y-bearing eggs","Diploid eggs"],"B","Human females are XX and all eggs carry X.","Human sex determination"),
("Human males are heterogametic because they produce:",["Only X-bearing sperm","Only Y-bearing sperm","X-bearing and Y-bearing sperm","Sperm without sex chromosomes"],"C","An XY male produces two types of sperm.","Human sex determination"),
("The sex of a human child is determined by:",["The ovum only","The type of sperm fertilising the ovum","Maternal nutrition","Mitochondria"],"B","The egg always contributes X, while sperm contributes X or Y.","Human sex determination"),
("In humans, an XX zygote develops as:",["Male","Female","Either sex randomly","Sterile male"],"B","XX is the typical female chromosomal constitution.","Human sex determination"),
("In birds, the heterogametic sex is:",["Male (ZZ)","Female (ZW)","Both sexes","Neither sex"],"B","Female birds produce Z- and W-bearing eggs.","Sex determination in birds"),
("In grasshoppers with XO type sex determination, males have:",["Two X chromosomes","One X chromosome and no second sex chromosome","XY chromosomes","ZW chromosomes"],"B","Males are XO and females are XX.","XO sex determination"),
("In honey bees, males develop from:",["Fertilised diploid eggs","Unfertilised haploid eggs","Somatic cells","Triploid eggs"],"B","Arrhenotokous parthenogenesis produces haploid males.","Haplodiploidy"),
("A male honey bee has:",["Diploid chromosome number","Haploid chromosome number","Two X chromosomes","One X and one Y"],"B","Drones arise from unfertilised eggs and are haploid.","Haplodiploidy"),
("A pedigree is particularly useful for studying inheritance in humans because:",["Controlled crosses are ethical","Human generation time is short","Controlled crosses are not possible and family records can be analysed","Humans have only one chromosome"],"C","Pedigrees trace trait transmission through generations.","Pedigree analysis"),
("In a pedigree, a square generally represents:",["Female","Male","Affected individual only","Carrier only"],"B","Squares denote males and circles denote females.","Pedigree symbols"),
("A filled symbol in a pedigree denotes:",["Unaffected individual","Affected individual","Only a carrier","Deceased individual"],"B","Shading indicates expression of the trait being studied.","Pedigree symbols"),
("Unaffected parents producing an affected child of either sex suggests most strongly:",["Autosomal dominant inheritance","Autosomal recessive inheritance","Y-linked inheritance","Mitochondrial inheritance"],"B","Both unaffected parents can be heterozygous carriers.","Pedigree analysis"),
("A trait appearing in every generation and affecting both sexes similarly is often:",["Autosomal dominant","Autosomal recessive only","Y-linked","X-linked recessive only"],"A","Autosomal dominant traits commonly show vertical transmission.","Pedigree analysis"),
("A Y-linked trait is transmitted from:",["Affected father to all daughters","Affected father to all sons","Affected mother to all sons","Carrier mother to daughters only"],"B","Only sons inherit the father's Y chromosome.","Y-linked inheritance"),
("An X-linked recessive trait is more frequent in males because males:",["Have two X chromosomes","Are hemizygous for X-linked genes","Never inherit X from mother","Have no autosomes"],"B","A single recessive allele on the male X is expressed.","X-linked inheritance"),
("A haemophilic father and a normal homozygous mother will have:",["All haemophilic sons","All carrier daughters and normal sons","All haemophilic daughters","Half haemophilic sons"],"B","Daughters receive the affected paternal X and normal maternal X; sons receive paternal Y.","Haemophilia"),
("A carrier woman for haemophilia marries a normal man. The probability that a son is haemophilic is:",["0","1/4","1/2","1"],"C","Each son receives Y from father and has a 1/2 chance of receiving the affected maternal X.","Haemophilia"),
("For the same couple, the probability that a randomly born child is a haemophilic son is:",["1/4","1/2","3/4","1"],"A","Probability of male is 1/2 and affected maternal X is 1/2, giving 1/4.","Haemophilia probability"),
("A colour-blind man marries a homozygous normal woman. Their daughters will be:",["All colour-blind","All carriers","Half carriers","All homozygous normal"],"B","Each daughter receives the father's affected X and a normal maternal X.","Colour blindness"),
("A carrier woman for an X-linked recessive trait and an affected man can produce affected daughters with probability:",["0","1/4 of all children","1/2 of all children","All daughters"],"B","An affected daughter must be female and receive the affected allele from the carrier mother: 1/2 × 1/2.","X-linked probability"),
("Father-to-son transmission rules out:",["Autosomal inheritance","Y-linked inheritance","X-linked inheritance","Polygenic inheritance"],"C","A father gives his Y, not his X, to sons.","Pedigree logic"),
("A trait transmitted by an affected mother to all children but by an affected father to none is suggestive of:",["Y-linkage","Mitochondrial inheritance","Autosomal dominance","X-linked dominance"],"B","Mitochondria are usually maternally inherited.","Maternal inheritance"),
("In an X-linked dominant trait, an affected father with a normal mother transmits the trait to:",["All sons only","All daughters only","Half of sons only","No children"],"B","All daughters receive his affected X; sons receive his Y.","X-linked dominant inheritance"),
("Consanguineous marriage increases the chance of:",["Autosomal recessive disorders","Y-linked traits only","New chromosome numbers","Dominant alleles disappearing"],"A","Relatives are more likely to carry the same rare recessive allele.","Pedigree and recessive disorders"),
("A normal woman whose father was haemophilic must be:",["Homozygous normal","A carrier","Haemophilic","Unable to have sons"],"B","She necessarily inherited the affected X from her father but is normal, so she is heterozygous.","X-linked inheritance")
]
for x in data4: add(*x)

# 106-130 mutations and disorders
data5=[
("A sudden heritable change in genetic material is called:",["Linkage","Mutation","Dominance","Assortment"],"B","Mutations generate new genetic variation.","Mutation"),
("Hugo de Vries proposed the mutation theory based largely on studies of:",["Pisum sativum","Oenothera lamarckiana","Drosophila","Neurospora"],"B","He studied evening primrose and emphasised sudden variations.","Mutation theory"),
("A point mutation involves:",["Change in the entire chromosome set","Change in a single base pair or a small DNA region","Fusion of two species","Only chromosome duplication"],"B","Point mutations alter one nucleotide pair or a very small sequence.","Gene mutation"),
("Sickle-cell anaemia is caused by:",["Deletion of an entire chromosome","A point mutation in the beta-globin gene","Trisomy 21","Failure of X chromosome segregation only"],"B","A single base substitution changes glutamic acid to valine in beta-globin.","Sickle-cell anaemia"),
("In sickle-cell anaemia, glutamic acid is replaced by valine at the:",["First position of alpha chain","Sixth position of beta chain","Sixth position of alpha chain","Last position of beta chain"],"B","The classic substitution occurs at beta-globin position six.","Sickle-cell anaemia"),
("Sickle-cell anaemia follows which inheritance pattern?",["Autosomal recessive","Autosomal dominant","X-linked recessive","Y-linked"],"A","The disease phenotype generally appears in homozygous HbS individuals.","Sickle-cell anaemia"),
("A heterozygous individual HbA HbS is usually described as having:",["Sickle-cell disease only","Sickle-cell trait","Thalassaemia","Down syndrome"],"B","Heterozygotes usually carry the trait and may show symptoms under low oxygen.","Sickle-cell trait"),
("Phenylketonuria is inherited as:",["Autosomal recessive","Autosomal dominant","X-linked dominant","Mitochondrial"],"A","It results from recessive mutation affecting phenylalanine metabolism.","Phenylketonuria"),
("Phenylketonuria involves impaired conversion of phenylalanine to:",["Tyrosine","Tryptophan","Glycine","Alanine"],"A","Deficient phenylalanine hydroxylase prevents normal tyrosine formation.","Phenylketonuria"),
("Thalassaemia is primarily a disorder of:",["Globin chain synthesis","Iron absorption only","Chromosome 21 number","Clotting factor VIII"],"A","Reduced synthesis of alpha or beta globin chains causes thalassaemia.","Thalassaemia"),
("Thalassaemia differs from sickle-cell anaemia because thalassaemia is mainly a:",["Qualitative structural defect only","Quantitative defect in globin synthesis","Chromosomal trisomy","Sex-linked disorder"],"B","Thalassaemia reduces the amount of a globin chain, whereas sickling changes beta-globin structure.","Thalassaemia"),
("Haemophilia is commonly inherited as:",["Autosomal dominant","X-linked recessive","Y-linked","Mitochondrial"],"B","The mutant clotting-factor allele is located on the X chromosome.","Haemophilia"),
("Colour blindness in humans is most commonly:",["Autosomal recessive","X-linked recessive","Y-linked dominant","Mitochondrial"],"B","Red-green colour blindness is a classic X-linked recessive trait.","Colour blindness"),
("Down syndrome is usually caused by:",["Monosomy X","Trisomy 21","XXY condition","Deletion of beta-globin gene"],"B","Nondisjunction can produce an extra chromosome 21.","Down syndrome"),
("The total chromosome number in a typical person with Down syndrome is:",["44","45","46","47"],"D","Trisomy 21 adds one chromosome to the normal diploid number.","Down syndrome"),
("Klinefelter syndrome usually has the karyotype:",["45,XO","47,XXY","47,XYY","47,XXX"],"B","The affected male carries an additional X chromosome.","Klinefelter syndrome"),
("A common feature of Klinefelter syndrome is:",["Female with rudimentary ovaries","Male with underdeveloped testes and possible gynaecomastia","Normal fertile male with 45 chromosomes","Trisomy 21"],"B","XXY individuals are male but often sterile with some feminised traits.","Klinefelter syndrome"),
("Turner syndrome usually has the karyotype:",["44,YO","45,XO","46,XX","47,XXY"],"B","Turner syndrome results from monosomy of the X chromosome.","Turner syndrome"),
("A typical person with Turner syndrome is:",["Male with XXY","Female with one X chromosome","Male with trisomy 21","Female with three X chromosomes"],"B","XO individuals develop as females with characteristic abnormalities.","Turner syndrome"),
("Aneuploidy results from:",["Addition or loss of one or a few chromosomes","Duplication of the entire genome only","A single base substitution only","Crossing over"],"A","Aneuploid cells have chromosome numbers not exact multiples of the haploid set.","Aneuploidy"),
("Polyploidy means:",["Loss of one chromosome","Gain of one chromosome","Presence of more than two complete chromosome sets","Mutation in one gene"],"C","Polyploids contain additional complete sets of chromosomes.","Polyploidy"),
("Nondisjunction is the failure of:",["DNA to replicate","Chromosomes or chromatids to separate properly","Genes to mutate","Gametes to fuse"],"B","Improper segregation can generate aneuploid gametes.","Nondisjunction"),
("A gamete with n+1 chromosomes fusing with a normal gamete produces a:",["Monosomic zygote","Trisomic zygote","Haploid zygote","Tetraploid zygote"],"B","The zygote has 2n+1 chromosomes.","Aneuploidy"),
("A gamete with n-1 chromosomes fusing with a normal gamete produces a:",["Monosomic zygote","Trisomic zygote","Triploid zygote","Normal zygote"],"A","The resulting chromosome number is 2n-1.","Aneuploidy"),
("Which disorder is correctly matched with its chromosomal basis?",["Turner syndrome - trisomy 21","Klinefelter syndrome - XXY","Down syndrome - XO","Haemophilia - trisomy 18"],"B","Klinefelter syndrome is associated with 47,XXY.","Chromosomal disorders")
]
for x in data5: add(*x)

# 131-150 Assertion Reason
aropts=[
"Both Assertion and Reason are true, and Reason is the correct explanation of Assertion",
"Both Assertion and Reason are true, but Reason is not the correct explanation of Assertion",
"Assertion is true, but Reason is false",
"Assertion is false, but Reason is true"
]
ars=[
("Assertion: The recessive trait reappears in the F2 generation. Reason: The two alleles of a gene segregate during gamete formation.","A","Segregation allows recessive alleles to reunite in homozygous offspring.","Law of segregation"),
("Assertion: A test cross can distinguish TT from Tt. Reason: The tester is homozygous recessive and reveals the gametes of the dominant-phenotype parent.","A","Different gametes from TT and Tt produce different progeny patterns.","Test cross"),
("Assertion: Incomplete dominance violates Mendel's law of segregation. Reason: The F2 genotypic ratio remains 1:2:1.","D","The assertion is false; alleles still segregate normally.","Incomplete dominance"),
("Assertion: IA and IB are codominant. Reason: Both alleles produce their respective antigens in an AB individual.","A","Simultaneous antigen expression demonstrates codominance.","ABO blood group"),
("Assertion: An individual can carry three ABO alleles. Reason: The ABO locus has three alleles in the human population.","D","The population has three alleles, but a diploid individual carries only two.","Multiple alleles"),
("Assertion: Linked genes may fail to show independent assortment. Reason: They are physically located on the same chromosome.","A","Physical association makes their inheritance non-independent.","Linkage"),
("Assertion: Recombination frequency can be used to estimate gene distance. Reason: Crossing over is generally more likely between genes farther apart.","A","Greater physical separation usually permits more crossovers.","Gene mapping"),
("Assertion: Recombination frequency between two genes can exceed 50%. Reason: Multiple crossovers are possible.","D","Multiple crossovers occur, but observable recombination does not exceed 50%.","Recombination"),
("Assertion: Human males are heterogametic. Reason: They form two types of sperm with respect to sex chromosomes.","A","Males produce X-bearing and Y-bearing sperm.","Human sex determination"),
("Assertion: In birds, the male determines the sex of offspring. Reason: Male birds are homogametic ZZ.","D","The reason is true, but heterogametic ZW females determine sex.","Sex determination in birds"),
("Assertion: X-linked recessive traits are more common in males. Reason: Males possess only one X chromosome.","A","There is no second allele to mask a recessive mutation in males.","X-linked inheritance"),
("Assertion: A haemophilic father transmits haemophilia directly to his sons. Reason: Sons receive the Y chromosome from their father.","D","The reason is true and explains why direct father-to-son X-linked transmission does not occur.","Haemophilia"),
("Assertion: Consanguinity can increase autosomal recessive disease risk. Reason: Related individuals may share alleles inherited from a common ancestor.","A","Shared ancestry raises the chance both partners carry the same recessive allele.","Pedigree genetics"),
("Assertion: Sickle-cell anaemia is caused by a point mutation. Reason: A single base substitution changes one amino acid in beta-globin.","A","The substitution produces valine in place of glutamic acid.","Sickle-cell anaemia"),
("Assertion: Thalassaemia is a quantitative haemoglobin defect. Reason: Synthesis of one globin chain is reduced or absent.","A","The amount of a globin chain is affected.","Thalassaemia"),
("Assertion: Down syndrome is caused by monosomy 21. Reason: Nondisjunction can alter chromosome number.","D","Nondisjunction is relevant, but Down syndrome is trisomy 21.","Down syndrome"),
("Assertion: Turner syndrome occurs in phenotypic females. Reason: The typical karyotype is 45,XO.","A","A single X chromosome supports female development but causes Turner features.","Turner syndrome"),
("Assertion: Klinefelter syndrome individuals are genetically male. Reason: They usually possess a Y chromosome along with two X chromosomes.","A","Presence of Y initiates male development in 47,XXY individuals.","Klinefelter syndrome"),
("Assertion: Polygenic traits often show continuous variation. Reason: Multiple genes contribute cumulative effects to the phenotype.","A","Many additive combinations create a spectrum of phenotypes.","Polygenic inheritance"),
("Assertion: Pleiotropy and polygenic inheritance mean the same thing. Reason: Pleiotropy is one gene affecting many traits, whereas polygenic inheritance is many genes affecting one trait.","D","The assertion is false while the reason correctly distinguishes them.","Pleiotropy and polygenic inheritance")
]
for a,b,c,d in ars: add(a,aropts,b,c,d,"NEET","Assertion-Reason")

# 151-160 Match the following
matches=[
("Match the inheritance pattern with the example: I. Incomplete dominance; II. Codominance; III. Polygenic inheritance; IV. Pleiotropy. a. ABO AB phenotype; b. Snapdragon flower colour; c. Human skin colour; d. Phenylketonuria",
 ["I-b, II-a, III-c, IV-d","I-a, II-b, III-d, IV-c","I-c, II-d, III-a, IV-b","I-d, II-c, III-b, IV-a"],"A","Snapdragon, AB blood group, skin colour and phenylketonuria are the respective examples.","Inheritance patterns"),
("Match the cross with the expected ratio: I. Tt × Tt phenotype; II. Tt × tt phenotype; III. Rr × Rr in incomplete dominance; IV. AaBb × aabb with independent assortment. a. 1:1; b. 3:1; c. 1:2:1; d. 1:1:1:1",
 ["I-b, II-a, III-c, IV-d","I-a, II-b, III-d, IV-c","I-c, II-d, III-a, IV-b","I-d, II-c, III-b, IV-a"],"A","The standard mono-, test-, incomplete-dominance and dihybrid-test-cross ratios match in that order.","Genetic ratios"),
("Match the scientist with contribution: I. Mendel; II. Sutton and Boveri; III. Morgan; IV. Sturtevant. a. Gene mapping; b. Laws of inheritance; c. Chromosomal theory; d. Linkage studies in Drosophila",
 ["I-b, II-c, III-d, IV-a","I-c, II-b, III-a, IV-d","I-d, II-a, III-c, IV-b","I-a, II-d, III-b, IV-c"],"A","These are the standard historical contributions described in NCERT.","History of genetics"),
("Match the condition with karyotype/basis: I. Down syndrome; II. Turner syndrome; III. Klinefelter syndrome; IV. Sickle-cell anaemia. a. 45,XO; b. Point mutation in beta-globin; c. Trisomy 21; d. 47,XXY",
 ["I-c, II-a, III-d, IV-b","I-a, II-c, III-b, IV-d","I-d, II-b, III-c, IV-a","I-b, II-d, III-a, IV-c"],"A","Each disorder is paired with its defining chromosomal or molecular basis.","Genetic disorders"),
("Match the organism with sex-determination system: I. Humans; II. Birds; III. Grasshopper; IV. Honey bee. a. XO male; b. XY male; c. ZW female; d. Haplodiploidy",
 ["I-b, II-c, III-a, IV-d","I-c, II-b, III-d, IV-a","I-a, II-d, III-b, IV-c","I-d, II-a, III-c, IV-b"],"A","Humans are XY, birds ZW, grasshoppers XO and honey bees haplodiploid.","Sex determination"),
("Match the pedigree clue with likely inheritance: I. Father to all sons; II. Affected father to all daughters but no sons; III. Unaffected parents with affected child of either sex; IV. Affected mother to all children but affected father to none. a. Autosomal recessive; b. Y-linked; c. Mitochondrial; d. X-linked dominant",
 ["I-b, II-d, III-a, IV-c","I-d, II-b, III-c, IV-a","I-a, II-c, III-d, IV-b","I-c, II-a, III-b, IV-d"],"A","These transmission patterns are diagnostic pedigree clues.","Pedigree analysis"),
("Match the term with definition: I. Linkage; II. Recombination; III. Centimorgan; IV. Nondisjunction. a. Failure of chromosome separation; b. Tendency of same-chromosome genes to be inherited together; c. New allele combinations; d. Unit corresponding to 1% recombination",
 ["I-b, II-c, III-d, IV-a","I-c, II-b, III-a, IV-d","I-d, II-a, III-c, IV-b","I-a, II-d, III-b, IV-c"],"A","The terms match their standard genetic definitions.","Genetic terminology"),
("Match the genotype with phenotype: I. IAIB; II. ii; III. HbAHbS; IV. HbSHbS. a. Sickle-cell trait; b. Blood group O; c. Blood group AB; d. Sickle-cell disease",
 ["I-c, II-b, III-a, IV-d","I-b, II-c, III-d, IV-a","I-d, II-a, III-c, IV-b","I-a, II-d, III-b, IV-c"],"A","The genotype-phenotype associations are direct.","Genotypes and phenotypes"),
("Match the meiotic event with stage: I. Synapsis; II. Crossing over; III. Chiasmata visible; IV. Homologous chromosomes separate. a. Diplotene; b. Pachytene; c. Anaphase I; d. Zygotene",
 ["I-d, II-b, III-a, IV-c","I-b, II-d, III-c, IV-a","I-a, II-c, III-d, IV-b","I-c, II-a, III-b, IV-d"],"A","Synapsis, crossing over, visible chiasmata and homolog separation occur in zygotene, pachytene, diplotene and anaphase I.","Meiosis and inheritance"),
("Match the concept with numerical expression: I. Gamete types from n heterozygous loci; II. Phenotypic classes for n additive polygenes; III. Homozygous recessive at two loci in AaBb × AaBb; IV. Recombinant percentage. a. 2n+1; b. Recombinants/total ×100; c. 2^n; d. 1/16",
 ["I-c, II-a, III-d, IV-b","I-a, II-c, III-b, IV-d","I-d, II-b, III-a, IV-c","I-b, II-d, III-c, IV-a"],"A","These formulas apply to independent heterozygous loci, additive classes, dihybrid probability and mapping.","Genetic calculations")
]
for x in matches: add(*x, level="NEET", typ="Match the Following")

assert len(qs)==160, len(qs)

doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(0.55); sec.bottom_margin=Inches(0.55)
sec.left_margin=Inches(0.62); sec.right_margin=Inches(0.62)
doc.styles["Normal"].font.name="Arial"; doc.styles["Normal"].font.size=Pt(9.3)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Principles of Inheritance and Variation"); r.bold=True; r.font.size=Pt(21)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("160 Original NCERT-Aligned NEET Questions | Medium-Hard | PYQ-Inspired"); r.italic=True; r.font.size=Pt(10)

p=doc.add_paragraph()
p.add_run("Composition: ").bold=True
p.add_run("130 single-correct MCQs, 20 assertion-reason questions and 10 match-the-following questions. Questions are original and use recurring NEET patterns without copying PYQs.")

p=doc.add_paragraph()
p.add_run("Syllabus coverage: ").bold=True
p.add_run("Mendelian inheritance, monohybrid and dihybrid crosses, test cross, deviations from Mendelism, ABO blood groups, polygenic inheritance, pleiotropy, chromosomal theory, linkage and recombination, sex determination, pedigree analysis, mutations and genetic disorders.")

p=doc.add_paragraph()
p.add_run("NCERT reference note: ").bold=True
p.add_run("Topic labels are used rather than page numbers because pagination differs across NCERT editions and rationalised printings.")

for i,item in enumerate(qs,1):
    para=doc.add_paragraph()
    para.paragraph_format.space_before=Pt(5.5); para.paragraph_format.space_after=Pt(1)
    rr=para.add_run(f"{i}. {item['q']}"); rr.bold=True
    for j,opt in enumerate(item["opts"]):
        op=doc.add_paragraph()
        op.paragraph_format.left_indent=Inches(0.2); op.paragraph_format.space_after=Pt(0)
        op.add_run(f"{chr(65+j)}. {opt}")
    meta=doc.add_paragraph()
    meta.paragraph_format.left_indent=Inches(0.2); meta.paragraph_format.space_after=Pt(0)
    a=meta.add_run(f"Answer: {item['ans']}  |  "); a.bold=True
    meta.add_run(f"Type: {item['type']}  |  Level: {item['level']}  |  NCERT topic: {item['topic']}")
    ex=doc.add_paragraph()
    ex.paragraph_format.left_indent=Inches(0.2); ex.paragraph_format.space_after=Pt(2)
    e=ex.add_run("Explanation: "); e.bold=True
    ex.add_run(item["exp"])
    if i<len(qs):
        line=doc.add_paragraph()
        pPr=line._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); bottom=OxmlElement("w:bottom")
        bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"3"); bottom.set(qn("w:space"),"1"); bottom.set(qn("w:color"),"DDDDDD")
        pBdr.append(bottom); pPr.append(pBdr)

for s in doc.sections:
    f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
    f.add_run("Medical Mania Biology Question Bank - Principles of Inheritance and Variation")

out=Path("/mnt/data/principles_of_inheritance_and_variation_160_questions.docx")
doc.save(out)
print(out, len(qs))

</USER_REQUEST>
<ADDITIONAL_METADATA>
The current local time is: 2026-07-21T12:07:43+05:30.
</ADDITIONAL_METADATA>